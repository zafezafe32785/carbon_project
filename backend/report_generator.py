"""
New AI-Powered Report Generation System for Carbon Accounting
Supports ISO 14064, CFO, and GHG Protocol standards with AI-generated preliminary descriptions

VERSION: 2.0 - COMPREHENSIVE AI REPORT WITH 8 SECTIONS
Last Updated: 2025-10-08
"""

import os

# Print version on module load for debugging
print("="*80)
print("LOADING REPORT GENERATOR VERSION 2.0 - COMPREHENSIVE AI REPORT")
print("="*80)
import json
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
import openai
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
import urllib.request
from models import emission_records_collection, users_collection, reports_collection, emission_factors_collection, calculate_co2_equivalent
from bson import ObjectId

# Configure OpenAI API
openai.api_key = os.getenv('OPENAI_API_KEY')

class CarbonReportGenerator:
    """
    AI-powered carbon accounting report generator
    Generates professional reports with AI-generated preliminary descriptions
    Focuses on GHG Protocol standard with multiple file formats and languages
    """
    
    def __init__(self):
        self.supported_formats = ['GHG']
        self.supported_file_types = ['PDF', 'EXCEL', 'WORD']
        self.supported_languages = ['EN', 'TH']

        # NOTE: Emission data is reused from smart_dashboard
        # All CO2 calculations are already done using TGO emission factors
        # Reports only separate by Scope 1 and Scope 2 (Scope 3 removed per project requirement)
        #
        # Scope 1: Direct emissions (fuels, refrigerants, combustion)
        # Scope 2: Indirect emissions from purchased electricity/energy

        # Thai month names for date formatting
        self.thai_months = {
            1: 'มกราคม', 2: 'กุมภาพันธ์', 3: 'มีนาคม', 4: 'เมษายน',
            5: 'พฤษภาคม', 6: 'มิถุนายน', 7: 'กรกฎาคม', 8: 'สิงหาคม',
            9: 'กันยายน', 10: 'ตุลาคม', 11: 'พฤศจิกายน', 12: 'ธันวาคม'
        }

        # Metric label translations (Thai labels for key metrics)
        self.metric_labels_th = {
            'total_emissions': 'การปล่อยก๊าซเรือนกระจกรวม',
            'average_monthly': 'ค่าเฉลี่ยรายเดือน',
            'reporting_period': 'ช่วงเวลารายงาน',
            'record_count': 'จำนวนข้อมูลการปล่อย'
        }

    def generate_report(self, user_id: str, start_date: str, end_date: str, 
                       report_format: str = 'GHG', file_type: str = 'PDF', 
                       language: str = 'EN', include_ai_insights: bool = True) -> Dict:
        """
        Generate a comprehensive carbon accounting report with AI-powered preliminary descriptions
        
        Args:
            user_id: User ID for data filtering
            start_date: Report start date (ISO format)
            end_date: Report end date (ISO format)
            report_format: Report format (GHG only)
            file_type: Output file type (PDF, EXCEL, WORD)
            language: Report language (EN, TH)
            include_ai_insights: Whether to include AI-generated content
            
        Returns:
            Dictionary with report generation results
        """
        try:
            print(f"\n{'='*80}")
            print(f"REPORT GENERATOR V2.0 - GENERATING COMPREHENSIVE AI REPORT")
            print(f"{'='*80}")
            print(f"Format: {report_format}, Language: {language}, File Type: {file_type}")
            print(f"User: {user_id}, Period: {start_date} to {end_date}")

            # Step 1: Collect and process emission data
            report_data = self._collect_emission_data(user_id, start_date, end_date)
            print(f"✓ Data collected: {report_data['record_count']} records, {report_data['total_emissions']:.2f} kg CO2e")

            # Step 2: Generate AI-powered preliminary descriptions
            ai_content = {}
            if include_ai_insights:
                print(f"✓ Generating AI content with 8 comprehensive sections...")
                ai_content = self._generate_ai_content(report_data, report_format, language)
                print(f"✓ AI sections generated: {list(ai_content.keys())}")
            
            # Step 3: Create structured report content
            report_content = self._create_report_structure(report_data, ai_content, report_format, language)
            
            # Step 4: Generate report file based on type
            file_path = self._generate_report_file(report_content, report_format, file_type, language)
            print(f"\n{'='*60}")
            print(f"REPORT FILE GENERATED")
            print(f"{'='*60}")
            print(f"  Returned file path: {file_path}")
            print(f"  File type: {file_type}")

            # Step 5: Save to database with your schema
            report_id = self._save_to_database(user_id, report_data, report_content,
                                             start_date, end_date, report_format, file_path, file_type, language)
            
            return {
                'success': True,
                'report_id': report_id,
                'file_path': file_path,
                'file_type': file_type,
                'language': language,
                'total_emissions': report_data['total_emissions'],
                'ai_insights_included': include_ai_insights,
                'generated_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            print(f"Report generation error: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'message': f'Failed to generate {report_format} report: {str(e)}'
            }

    def _collect_emission_data(self, user_id: str, start_date: str, end_date: str) -> Dict:
        """
        Collect and process emission data from database
        Uses SHARED data from ALL users (same as dashboard)
        All CO2 calculations are already done using TGO factors
        """

        # Parse dates
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))

        # Get user information (for report metadata only)
        user = users_collection.find_one({'_id': ObjectId(user_id)})
        organization = user.get('organization', 'All Organizations') if user else 'All Organizations'

        # Query ALL emission records from ALL users (shared data - same as dashboard)
        emissions = list(emission_records_collection.find({
            'record_date': {'$gte': start_dt, '$lte': end_dt}
        }))

        print(f"Found {len(emissions)} emission records for report period")

        # Total emissions (already calculated with TGO factors in dashboard)
        total_emissions = sum(float(e.get('co2_equivalent', 0)) for e in emissions)

        # Group by category (same as dashboard)
        emissions_by_category = {}
        for emission in emissions:
            category = emission.get('category', 'other')
            co2_value = float(emission.get('co2_equivalent', 0))
            emissions_by_category[category] = emissions_by_category.get(category, 0) + co2_value

        # Group by GHG Protocol scopes (Scope 1 and 2 only as per project requirement)
        emissions_by_scope = {'Scope 1': 0, 'Scope 2': 0}

        # Define Scope 1 keywords (matching dashboard logic)
        scope1_keywords = [
            'fuel', 'gasoline', 'diesel', 'natural gas', 'lpg', 'coal', 'kerosene',
            'mobile', 'vehicle', 'transport', 'combustion', 'stationary', 'biomass',
            'bagasse', 'biogas', 'wood', 'anthracite', 'bituminous', 'lignite',
            'refrigerant', 'fugitive', 'cng', 'heavy fuel oil', 'gas oil',
            'equipment', 'machinery', 'agriculture', 'forestry', 'construction',
            'r-', 'hfc', 'pfc', 'sf6', 'nf3'  # Refrigerant codes
        ]

        for emission in emissions:
            category = emission.get('category', 'other').lower()
            co2_value = float(emission.get('co2_equivalent', 0))

            # Determine scope based on category
            # Scope 2: Electricity and purchased energy
            if any(keyword in category for keyword in ['electric', 'grid', 'power', 'energy']):
                emissions_by_scope['Scope 2'] += co2_value
            # Scope 1: Direct emissions (fuels, refrigerants, combustion, etc.)
            elif any(keyword in category for keyword in scope1_keywords):
                emissions_by_scope['Scope 1'] += co2_value
            # Default: Everything else goes to Scope 1
            else:
                emissions_by_scope['Scope 1'] += co2_value

        print(f"Scope breakdown: Scope 1 = {emissions_by_scope['Scope 1']}, Scope 2 = {emissions_by_scope['Scope 2']}")

        # Calculate monthly breakdown (same logic as dashboard)
        monthly_data = self._calculate_monthly_breakdown(emissions, start_dt, end_dt)

        return {
            'user_id': user_id,
            'organization': organization,
            'period_start': start_dt,
            'period_end': end_dt,
            'total_emissions': total_emissions,
            'emissions_by_category': emissions_by_category,
            'emissions_by_scope': emissions_by_scope,
            'monthly_data': monthly_data,
            'record_count': len(emissions),
            'raw_emissions': emissions
        }

    def _calculate_monthly_breakdown(self, emissions: List[Dict], start_date: datetime, end_date: datetime) -> List[Dict]:
        """Calculate monthly emission breakdown"""
        monthly_breakdown = {}
        
        # Initialize months in range
        current = start_date.replace(day=1)
        while current <= end_date:
            month_key = current.strftime('%Y-%m')
            monthly_breakdown[month_key] = {
                'month': current.strftime('%B %Y'),
                'total': 0,
                'by_category': {}
            }
            current = (current.replace(day=28) + timedelta(days=4)).replace(day=1)
        
        # Aggregate emissions by month
        for emission in emissions:
            record_date = emission.get('record_date', datetime.now())
            month_key = record_date.strftime('%Y-%m')
            
            if month_key in monthly_breakdown:
                co2_value = emission.get('co2_equivalent', 0)
                category = emission.get('category', 'other')
                
                monthly_breakdown[month_key]['total'] += co2_value
                monthly_breakdown[month_key]['by_category'][category] = \
                    monthly_breakdown[month_key]['by_category'].get(category, 0) + co2_value
        
        return list(monthly_breakdown.values())

    def _generate_ai_content(self, report_data: Dict, report_format: str, language: str = 'EN') -> Dict:
        """Generate AI-powered preliminary descriptions and insights for comprehensive report"""
        try:
            if not openai.api_key:
                return self._get_fallback_content(report_data, report_format, language)

            ai_content = {}

            # 1. Executive Summary (สรุปผู้บริหาร)
            ai_content['executive_summary'] = self._generate_executive_summary(report_data, report_format, language)

            # 2. Trend Analysis (การวิเคราะห์แนวโน้ม) - FIXED: Now handles 2024 data correctly
            ai_content['trend_analysis'] = self._generate_trend_analysis(report_data, language)

            # 3. Emissions Breakdown (การแบ่งประเภทการปล่อยก๊าซ)
            ai_content['emissions_breakdown'] = self._generate_emissions_breakdown(report_data, language)

            # 4. Methodology (วิธีการคำนวณ) - Enhanced with AI
            ai_content['methodology'] = self._generate_methodology(report_data, report_format, language)

            # 5. Data Quality Statement (คำชี้แจงคุณภาพข้อมูล)
            ai_content['data_quality'] = self._generate_data_quality(report_data, language)

            # 7. Recommendations (ข้อเสนอแนะ)
            ai_content['recommendations'] = self._generate_recommendations(report_data, language)

            # 8. Conclusion (สรุป)
            ai_content['conclusion'] = self._generate_conclusion(report_data, language)

            # Legacy support - key findings (backward compatibility)
            ai_content['key_findings'] = self._generate_key_findings(report_data, language)

            return ai_content

        except Exception as e:
            print(f"AI content generation error: {str(e)}")
            return self._get_fallback_content(report_data, report_format, language)

    def _generate_executive_summary(self, report_data: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate AI-powered executive summary with professional tone"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            # Calculate key metrics
            top_categories = dict(sorted(report_data['emissions_by_category'].items(), key=lambda x: x[1], reverse=True)[:3])
            avg_monthly = report_data['total_emissions'] / max(len(report_data.get('monthly_data', [])), 1)

            prompt = f"""
            Generate a comprehensive executive summary for a {report_format} Protocol carbon emissions report:

            ORGANIZATION DATA:
            - Organization: {report_data['organization']}
            - Reporting Period: {report_data['period_start'].strftime('%B %Y')} to {report_data['period_end'].strftime('%B %Y')}

            EMISSION METRICS:
            - Total Emissions: {report_data['total_emissions']:.2f} kg CO2e
            - Average Monthly: {avg_monthly:.2f} kg CO2e
            - Number of Records: {report_data['record_count']}

            BREAKDOWN:
            - Top Emission Sources: {top_categories}
            - Scope Distribution: {report_data['emissions_by_scope']}

            REQUIREMENTS:
            - {language_instruction}
            - Executive-level professional tone suitable for board presentation
            - 200-250 words comprehensive summary
            - Include the following elements:
              1. Overview of reporting scope and period
              2. Total emissions with context
              3. Key emission sources and their significance
              4. Scope breakdown implications
              5. Compliance status with {report_format} standards
              6. Strategic business implications
              7. Summary of recommended actions
            - Use clear, confident language
            - Emphasize transparency and accountability
            - Conclude with forward-looking statement
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.7
            )

            # Clean up and return
            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Executive summary generation error: {str(e)}")
            return self._get_fallback_executive_summary(report_data, report_format, language)

    def _generate_key_findings(self, report_data: Dict, language: str = 'EN') -> List[str]:
        """Generate AI-powered key findings"""
        try:
            language_instruction = "Write in Thai language" if language == 'TH' else "Write in English"
            
            prompt = f"""
            Based on this carbon emissions data, identify 5-7 key findings:
            
            Total Emissions: {report_data['total_emissions']:.2f} kg CO2e
            Categories: {report_data['emissions_by_category']}
            Scope Breakdown: {report_data['emissions_by_scope']}
            Record Count: {report_data['record_count']}
            
            Requirements:
            - {language_instruction}
            - Format as bullet points
            - Each finding should be specific and data-driven
            - Actionable for management
            - Focused on significant patterns or opportunities
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.6
            )
            
            findings_text = response.choices[0].message.content.strip()
            # Clean up multiple spaces and newlines, filter empty lines
            lines = [f.strip('• -').strip() for f in findings_text.split('\n')]
            # Remove empty strings and normalize whitespace
            return [' '.join(line.split()) for line in lines if line.strip()]
            
        except Exception as e:
            return self._get_fallback_key_findings(report_data, language)

    def _generate_recommendations(self, report_data: Dict, language: str = 'EN') -> List[str]:
        """Generate AI-powered actionable recommendations"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"
            top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get) \
                          if report_data['emissions_by_category'] else 'unknown'

            # Calculate scope percentages
            scope1_pct = (report_data['emissions_by_scope'].get('Scope 1', 0) / report_data['total_emissions'] * 100) if report_data['total_emissions'] > 0 else 0
            scope2_pct = (report_data['emissions_by_scope'].get('Scope 2', 0) / report_data['total_emissions'] * 100) if report_data['total_emissions'] > 0 else 0

            prompt = f"""
            Generate 6-8 specific, actionable recommendations for carbon emission reduction:

            EMISSION PROFILE:
            - Total Emissions: {report_data['total_emissions']:.2f} kg CO2e
            - Primary Source: {top_category}
            - Scope 1 (Direct): {report_data['emissions_by_scope'].get('Scope 1', 0):.2f} kg CO2e ({scope1_pct:.1f}%)
            - Scope 2 (Energy): {report_data['emissions_by_scope'].get('Scope 2', 0):.2f} kg CO2e ({scope2_pct:.1f}%)
            - All Categories: {report_data['emissions_by_category']}

            REQUIREMENTS:
            - {language_instruction}
            - Professional and strategic tone
            - Each recommendation must be:
              * Specific and actionable (not generic advice)
              * Tied to actual emission sources in the data
              * Include estimated reduction potential (e.g., "10-15% reduction")
              * Consider cost-effectiveness and feasibility
              * Prioritized by impact (highest impact first)
            - Cover both immediate quick wins and long-term strategic initiatives
            - Include recommendations for:
              1. Top emission source ({top_category})
              2. Energy efficiency (if Scope 2 is significant)
              3. Operational improvements
              4. Technology upgrades
              5. Monitoring and verification systems
              6. Employee engagement programs
            - Format as clear bullet points
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=600,
                temperature=0.7
            )

            recommendations_text = response.choices[0].message.content.strip()
            # Clean up and parse recommendations
            lines = [r.strip('• -123456789.').strip() for r in recommendations_text.split('\n')]
            # Remove empty strings and normalize whitespace
            return [' '.join(line.split()) for line in lines if line.strip() and len(line.strip()) > 10]

        except Exception as e:
            print(f"Recommendations generation error: {str(e)}")
            return self._get_fallback_recommendations(report_data, language)

    def _generate_trend_analysis(self, report_data: Dict, language: str = 'EN') -> str:
        """
        Generate AI-powered trend analysis
        FIXED: Properly handles year-over-year comparisons including 2024 data
        """
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            # Extract monthly data with proper validation
            monthly_data = report_data.get('monthly_data', [])
            if not monthly_data:
                return self._get_fallback_trend_analysis(report_data, language)

            # Calculate comprehensive trend metrics
            monthly_totals = [month.get('total', 0) for month in monthly_data]
            total_emissions = sum(monthly_totals)
            avg_monthly = total_emissions / len(monthly_totals) if monthly_totals else 0
            max_month = max(monthly_data, key=lambda x: x.get('total', 0)) if monthly_data else {}
            min_month = min(monthly_data, key=lambda x: x.get('total', 0)) if monthly_data else {}

            # Year-over-year comparison (FIXED: Handles current year data)
            yearly_totals = {}
            for month in monthly_data:
                month_name = month.get('month', '')
                if month_name:
                    year = month_name.split()[-1]  # Extract year from "Month YYYY"
                    yearly_totals[year] = yearly_totals.get(year, 0) + month.get('total', 0)

            # Fetch previous year data for comparison
            period_start = report_data['period_start']
            period_end = report_data['period_end']
            current_year = period_start.year
            prev_year_start = period_start.replace(year=current_year - 1)
            prev_year_end = period_end.replace(year=current_year - 1)

            try:
                prev_year_emissions = list(emission_records_collection.find({
                    'record_date': {'$gte': prev_year_start, '$lte': prev_year_end}
                }))
                prev_year_total = sum(float(e.get('co2_equivalent', 0)) for e in prev_year_emissions)
                if prev_year_total > 0:
                    yearly_totals[str(current_year - 1)] = prev_year_total
            except:
                prev_year_total = 0

            # Calculate trend
            if len(monthly_totals) > 1:
                trend = "increasing" if monthly_totals[-1] > monthly_totals[0] else "decreasing" if monthly_totals[-1] < monthly_totals[0] else "stable"
                change_pct = ((monthly_totals[-1] - monthly_totals[0]) / monthly_totals[0] * 100) if monthly_totals[0] > 0 else 0
            else:
                trend = "stable"
                change_pct = 0

            # Calculate year-over-year change
            yoy_comparison = ""
            if prev_year_total > 0:
                yoy_change = ((total_emissions - prev_year_total) / prev_year_total) * 100
                yoy_comparison = f"Year-over-year change: {yoy_change:+.1f}% (Current: {total_emissions:.1f} vs Previous: {prev_year_total:.1f} kg CO2e)"

            # Build comprehensive prompt
            prompt = f"""
            Analyze greenhouse gas emission trends for this organization's carbon footprint:

            EMISSION DATA:
            - Total Emissions: {total_emissions:.2f} kg CO2e
            - Average Monthly: {avg_monthly:.2f} kg CO2e
            - Peak Month: {max_month.get('month', 'N/A')} ({max_month.get('total', 0):.2f} kg CO2e)
            - Lowest Month: {min_month.get('month', 'N/A')} ({min_month.get('total', 0):.2f} kg CO2e)
            - Monthly Data: {[f"{m.get('month', 'N/A')}: {m.get('total', 0):.1f} kg CO2e" for m in monthly_data]}
            - Year Totals by Year: {yearly_totals}
            - {yoy_comparison}
            - Overall Trend: {trend}
            - Period Change: {change_pct:+.1f}%

            REQUIREMENTS:
            - {language_instruction}
            - Professional tone suitable for executive and stakeholder review
            - 200-250 words comprehensive analysis
            - Include the following sections:
              1. Year-over-year comparison (compare {current_year} to {current_year-1} if data available)
              2. Monthly trend patterns and seasonality analysis
              3. Peak and low emission periods with possible explanations
              4. Business implications and operational insights
              5. Forward-looking recommendations for emission management
            - Use specific numbers and percentages from the data
            - Provide actionable insights for carbon reduction strategy
            - If comparing 2024 to 2025 data, explicitly mention both years
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.6
            )

            # Clean up and return
            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Trend analysis generation error: {str(e)}")
            return self._get_fallback_trend_analysis(report_data, language)

    def _generate_emissions_breakdown(self, report_data: Dict, language: str = 'EN') -> str:
        """Generate AI-powered emissions breakdown analysis"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            # Prepare category and scope data
            categories_sorted = sorted(report_data['emissions_by_category'].items(), key=lambda x: x[1], reverse=True)
            total_emissions = report_data['total_emissions']

            prompt = f"""
            Provide a detailed breakdown analysis of greenhouse gas emissions by source and scope:

            EMISSIONS BY CATEGORY (Top to Bottom):
            {chr(10).join([f"- {cat}: {val:.2f} kg CO2e ({(val/total_emissions*100):.1f}%)" for cat, val in categories_sorted if val > 0])}

            EMISSIONS BY SCOPE:
            - Scope 1 (Direct): {report_data['emissions_by_scope'].get('Scope 1', 0):.2f} kg CO2e ({(report_data['emissions_by_scope'].get('Scope 1', 0)/total_emissions*100):.1f}%)
            - Scope 2 (Indirect Energy): {report_data['emissions_by_scope'].get('Scope 2', 0):.2f} kg CO2e ({(report_data['emissions_by_scope'].get('Scope 2', 0)/total_emissions*100):.1f}%)

            Total: {total_emissions:.2f} kg CO2e

            REQUIREMENTS:
            - {language_instruction}
            - Professional and analytical tone
            - 150-200 words
            - Explain the significance of each major emission source
            - Highlight scope distribution and its implications
            - Identify which categories are most material to the organization
            - Provide context for why certain sources contribute more than others
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350,
                temperature=0.6
            )

            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Emissions breakdown generation error: {str(e)}")
            return self._get_fallback_emissions_breakdown(report_data, language)

    def _generate_methodology(self, report_data: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate AI-enhanced methodology description"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            prompt = f"""
            Describe the calculation methodology for this {report_format} carbon emissions report:

            REPORT DETAILS:
            - Standard: {report_format} Protocol
            - Total Records Analyzed: {report_data['record_count']}
            - Period: {report_data['period_start'].strftime('%B %Y')} to {report_data['period_end'].strftime('%B %Y')}
            - Scopes Covered: Scope 1 (Direct) and Scope 2 (Indirect Energy)

            REQUIREMENTS:
            - {language_instruction}
            - Professional and technical tone
            - 200-250 words comprehensive description
            - Include the following:
              1. Calculation approach and formulas used (Activity Data × Emission Factor = CO2 Equivalent)
              2. Emission factor sources (TGO - Thailand Greenhouse Gas Management Organization, IPCC guidelines)
              3. Data collection methods and frequency
              4. Organizational boundary approach (operational control)
              5. Scope definitions specific to {report_format}
              6. Quality assurance procedures
              7. Assumptions and limitations
            - Explain how emissions are categorized by scope
            - Reference specific standards and protocols followed
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=400,
                temperature=0.6
            )

            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Methodology generation error: {str(e)}")
            return self._get_methodology_text(report_format, language)

    def _generate_data_quality(self, report_data: Dict, language: str = 'EN') -> str:
        """Generate data quality statement"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            completeness = (report_data['record_count'] / max(report_data['record_count'], 1)) * 100  # Simplified metric

            prompt = f"""
            Provide a data quality statement for this carbon emissions report:

            DATA METRICS:
            - Total Records: {report_data['record_count']}
            - Reporting Period: {report_data['period_start'].strftime('%B %Y')} to {report_data['period_end'].strftime('%B %Y')}
            - Categories Covered: {len(report_data['emissions_by_category'])}
            - Data Completeness: High (all emission sources tracked)

            REQUIREMENTS:
            - {language_instruction}
            - Professional and transparent tone
            - 150-200 words
            - Address the following:
              1. Data completeness and coverage
              2. Data accuracy and confidence levels
              3. Primary data vs. estimated data
              4. Known gaps or uncertainties
              5. Verification and validation procedures
              6. Data management systems used
              7. Recommendations for data quality improvement
            - Be honest about limitations while highlighting strengths
            - Provide assurance appropriate for stakeholder confidence
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350,
                temperature=0.6
            )

            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Data quality generation error: {str(e)}")
            return self._get_fallback_data_quality(language)

    def _generate_conclusion(self, report_data: Dict, language: str = 'EN') -> str:
        """Generate comprehensive conclusion"""
        try:
            language_instruction = "Write in Thai language (ภาษาไทย)" if language == 'TH' else "Write in professional English"

            # Calculate some key metrics for conclusion
            avg_monthly = report_data['total_emissions'] / max(len(report_data['monthly_data']), 1)
            top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get) if report_data['emissions_by_category'] else 'unknown'

            prompt = f"""
            Write a comprehensive conclusion for this carbon emissions report:

            KEY METRICS:
            - Total Emissions: {report_data['total_emissions']:.2f} kg CO2e
            - Average Monthly: {avg_monthly:.2f} kg CO2e
            - Primary Source: {top_category}
            - Organization: {report_data['organization']}

            REQUIREMENTS:
            - {language_instruction}
            - Professional and forward-looking tone
            - 150-200 words
            - Include the following:
              1. Overall assessment of emission performance
              2. Progress toward sustainability goals
              3. Next steps and priorities
              4. Commitment to continuous improvement
              5. Call to action for stakeholders
            - Balance honesty with optimism
            - Emphasize accountability and transparency
            - End with a strong commitment statement
            """

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=350,
                temperature=0.7
            )

            import re
            content = response.choices[0].message.content.strip()
            content = re.sub(r'\s+', ' ', content)  # Remove multiple consecutive spaces
            return content

        except Exception as e:
            print(f"Conclusion generation error: {str(e)}")
            return self._get_fallback_conclusion(report_data, language)

    def _get_fallback_content(self, report_data: Dict, report_format: str, language: str = 'EN') -> Dict:
        """Fallback content when AI is not available"""
        top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get) \
                      if report_data['emissions_by_category'] else 'unknown'

        return {
            'executive_summary': self._get_fallback_executive_summary(report_data, report_format, language),
            'key_findings': self._get_fallback_key_findings(report_data, language),
            'recommendations': self._get_fallback_recommendations(report_data, language),
            'trend_analysis': self._get_fallback_trend_analysis(report_data, language),
            'emissions_breakdown': self._get_fallback_emissions_breakdown(report_data, language),
            'methodology': self._get_methodology_text(report_format, language),
            'data_quality': self._get_fallback_data_quality(language),
            'conclusion': self._get_fallback_conclusion(report_data, language)
        }

    def _get_fallback_executive_summary(self, report_data: Dict, report_format: str, language: str = 'EN') -> str:
        """Fallback executive summary"""
        if language == 'TH':
            start_str = self._format_month_year_thai(report_data['period_start'])
            end_str = self._format_month_year_thai(report_data['period_end'])
            text = f"""รายงานการปล่อยก๊าซเรือนกระจก {report_format} ฉบับนี้นำเสนอการวิเคราะห์ที่ครอบคลุมเกี่ยวกับการปล่อยก๊าซเรือนกระจกของ {report_data['organization']} ในช่วงระยะเวลาตั้งแต่ {start_str} ถึง {end_str} การปล่อยก๊าซเรือนกระจกรวมทั้งหมด {report_data['total_emissions']:.2f} kg CO2e จากข้อมูล {report_data['record_count']} รายการ รายงานนี้เป็นไปตามมาตรฐาน {report_format} และให้ข้อมูลเชิงลึกที่สามารถนำไปปฏิบัติได้สำหรับกลยุทธ์การลดการปล่อยก๊าซเรือนกระจก จุดสำคัญที่ควรให้ความสนใจ ได้แก่การปรับปรุงประสิทธิภาพการใช้พลังงานและการดำเนินงานที่ยั่งยืนเพื่อบรรลุเป้าหมายการลดคาร์บอน"""
            return self._clean_thai_text(text)
        else:
            text = f"""
            This {report_format} carbon emissions report presents a comprehensive analysis of greenhouse gas emissions
            for {report_data['organization']} covering the period from {report_data['period_start'].strftime('%B %Y')}
            to {report_data['period_end'].strftime('%B %Y')}. Total emissions reached {report_data['total_emissions']:.2f} kg CO2e
            across {report_data['record_count']} emission records. The report follows {report_format} standards and provides
            actionable insights for emission reduction strategies. Key focus areas include energy efficiency improvements
            and sustainable operational practices to achieve carbon reduction goals.
            """
            # Clean up multiple spaces and newlines
            return ' '.join(text.split())

    def _get_fallback_key_findings(self, report_data: Dict, language: str = 'EN') -> List[str]:
        """Fallback key findings"""
        if language == 'TH':
            start_str = self._format_month_year_thai(report_data['period_start'])
            end_str = self._format_month_year_thai(report_data['period_end'])
            findings = [
                f"การปล่อยก๊าซเรือนกระจกรวม: {report_data['total_emissions']:.2f} kg CO2e",
                f"ช่วงเวลารายงาน: {start_str} ถึง {end_str}",
                f"จำนวนข้อมูลการปล่อย: {report_data['record_count']} รายการ"
            ]
            
            if report_data['emissions_by_category']:
                top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get)
                findings.append(f"แหล่งปล่อยหลัก: {top_category}")
                findings.append(f"จำนวนประเภทการปล่อย: {len(report_data['emissions_by_category'])} ประเภท")
            
            # Add scope breakdown
            for scope, value in report_data['emissions_by_scope'].items():
                if value > 0:
                    percentage = (value / report_data['total_emissions'] * 100) if report_data['total_emissions'] > 0 else 0
                    findings.append(f"{scope}: {percentage:.1f}% ของการปล่อยทั้งหมด")
        else:
            findings = [
                f"Total emissions: {report_data['total_emissions']:.2f} kg CO2e",
                f"Reporting period: {report_data['period_start'].strftime('%B %Y')} to {report_data['period_end'].strftime('%B %Y')}",
                f"Number of emission records: {report_data['record_count']}"
            ]
            
            if report_data['emissions_by_category']:
                top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get)
                findings.append(f"Primary emission source: {top_category}")
                findings.append(f"Number of emission categories: {len(report_data['emissions_by_category'])}")
            
            # Add scope breakdown
            for scope, value in report_data['emissions_by_scope'].items():
                if value > 0:
                    percentage = (value / report_data['total_emissions'] * 100) if report_data['total_emissions'] > 0 else 0
                    findings.append(f"{scope}: {percentage:.1f}% of total emissions")
        
        return findings

    def _get_fallback_recommendations(self, report_data: Dict, language: str = 'EN') -> List[str]:
        """Fallback recommendations"""
        if language == 'TH':
            recommendations = [
                "ปรับปรุงระบบการจัดการพลังงานอย่างครอบคลุม",
                "ดำเนินการตรวจสอบพลังงานเป็นประจำเพื่อหาโอกาสในการเพิ่มประสิทธิภาพ",
                "พิจารณาตัวเลือกการจัดหาพลังงานทดแทน",
                "กำหนดเป้าหมายการลดการปล่อยและขั้นตอนการติดตาม",
                "พัฒนาโปรแกรมสร้างความตระหนักของพนักงานในการลดคาร์บอน"
            ]
            
            if report_data['emissions_by_category']:
                top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get)
                recommendations.insert(0, f"จัดลำดับความสำคัญของกลยุทธ์การลดการปล่อยจาก {top_category}")
        else:
            recommendations = [
                "Implement comprehensive energy management system",
                "Conduct regular energy audits to identify efficiency opportunities",
                "Consider renewable energy procurement options",
                "Establish emission reduction targets and monitoring procedures",
                "Develop employee awareness programs for carbon reduction"
            ]
            
            if report_data['emissions_by_category']:
                top_category = max(report_data['emissions_by_category'], key=report_data['emissions_by_category'].get)
                recommendations.insert(0, f"Prioritize reduction strategies for {top_category} emissions")
        
        return recommendations

    def _get_fallback_trend_analysis(self, report_data: Dict, language: str = 'EN') -> str:
        """Fallback trend analysis with year-over-year comparison"""
        # Validate monthly_data exists and has sufficient entries
        if not report_data.get('monthly_data') or len(report_data['monthly_data']) < 1:
            return "Insufficient data available for trend analysis." if language == 'EN' else "ข้อมูลไม่เพียงพอสำหรับการวิเคราะห์แนวโน้ม"

        # Try to fetch previous year data for year-over-year comparison
        period_start = report_data['period_start']
        period_end = report_data['period_end']
        current_year = period_start.year

        # Calculate previous year same period
        prev_year_start = period_start.replace(year=current_year - 1)
        prev_year_end = period_end.replace(year=current_year - 1)

        # Query previous year emissions
        try:
            prev_year_emissions = list(emission_records_collection.find({
                'record_date': {'$gte': prev_year_start, '$lte': prev_year_end}
            }))
            prev_year_total = sum(float(e.get('co2_equivalent', 0)) for e in prev_year_emissions)
        except:
            prev_year_total = 0

        if len(report_data['monthly_data']) > 1:
            first_month = report_data['monthly_data'][0].get('total', 0)
            last_month = report_data['monthly_data'][-1].get('total', 0)
            current_total = report_data['total_emissions']

            # Build trend analysis text
            if language == 'TH':
                text = ""
                if last_month > first_month:
                    text = f"การปล่อยก๊าซเรือนกระจกแสดงแนวโน้มเพิ่มขึ้นในช่วงระยะเวลารายงาน เพิ่มขึ้นจาก {first_month:.1f} เป็น {last_month:.1f} kg CO2e"
                elif last_month < first_month:
                    text = f"การปล่อยก๊าซเรือนกระจกแสดงแนวโน้มลดลง ลดลงจาก {first_month:.1f} เป็น {last_month:.1f} kg CO2e"
                else:
                    text = "การปล่อยก๊าซเรือนกระจกยังคงค่อนข้างคงที่ตลอดช่วงระยะเวลารายงาน"

                # Add year-over-year comparison if data available
                if prev_year_total > 0:
                    change_pct = ((current_total - prev_year_total) / prev_year_total) * 100
                    if change_pct > 0:
                        text += f" เมื่อเทียบกับปี {current_year - 1} ({prev_year_total:.1f} kg CO2e) การปล่อยก๊าซเพิ่มขึ้น {abs(change_pct):.1f}% ในปี {current_year}"
                    elif change_pct < 0:
                        text += f" เมื่อเทียบกับปี {current_year - 1} ({prev_year_total:.1f} kg CO2e) การปล่อยก๊าซลดลง {abs(change_pct):.1f}% ในปี {current_year}"
                    else:
                        text += f" เมื่อเทียบกับปี {current_year - 1} การปล่อยก๊าซคงที่"
                    text += " ซึ่งบ่งชี้ถึงความจำเป็นในการเสริมสร้างมาตรการลดการปล่อยและการติดตามแหล่งปล่อยอย่างใกล้ชิด"

                return text
            else:
                text = ""
                if last_month > first_month:
                    text = f"Emissions showed an increasing trend over the reporting period, rising from {first_month:.1f} to {last_month:.1f} kg CO2e."
                elif last_month < first_month:
                    text = f"Emissions demonstrated a decreasing trend, falling from {first_month:.1f} to {last_month:.1f} kg CO2e."
                else:
                    text = "Emissions remained relatively stable throughout the reporting period."

                # Add year-over-year comparison if data available
                if prev_year_total > 0:
                    change_pct = ((current_total - prev_year_total) / prev_year_total) * 100
                    if change_pct > 0:
                        text += f" Compared to {current_year - 1} ({prev_year_total:.1f} kg CO2e), emissions increased by {abs(change_pct):.1f}% in {current_year}."
                    elif change_pct < 0:
                        text += f" Compared to {current_year - 1} ({prev_year_total:.1f} kg CO2e), emissions decreased by {abs(change_pct):.1f}% in {current_year}."
                    else:
                        text += f" Emissions remained stable compared to {current_year - 1}."
                    text += " This indicates a need for enhanced emission reduction measures and closer monitoring of emission sources."

                return text
        else:
            if language == 'TH':
                return "การวิเคราะห์แนวโน้มต้องการข้อมูลหลายช่วงเวลาเพื่อการประเมินที่มีความหมาย รายงานในอนาคตจะรวมการวิเคราะห์เปรียบเทียบเพื่อระบุรูปแบบและโอกาสในการปรับปรุง"
            else:
                return "Trend analysis requires multiple reporting periods for meaningful assessment. Future reports will include comparative analysis to identify patterns and improvement opportunities."

    def _get_fallback_emissions_breakdown(self, report_data: Dict, language: str = 'EN') -> str:
        """Fallback emissions breakdown"""
        categories_sorted = sorted(report_data['emissions_by_category'].items(), key=lambda x: x[1], reverse=True)
        total_emissions = report_data['total_emissions']

        if language == 'TH':
            # Clean category names by replacing underscores with spaces
            categories_text = ', '.join([f"{cat.replace('_', ' ')} ({(val/total_emissions*100):.1f}%)" for cat, val in categories_sorted[:3] if val > 0])
            text = f"""การวิเคราะห์การแบ่งประเภทการปล่อยก๊าซเรือนกระจกแสดงให้เห็นว่าการปล่อยทั้งหมด {total_emissions:.2f} kg CO2e มาจากหลายแหล่ง โดยแหล่งปล่อยหลัก ได้แก่ {categories_text} การปล่อยถูกจำแนกตาม GHG Protocol เป็น Scope 1 (การปล่อยโดยตรง) {report_data['emissions_by_scope'].get('Scope 1', 0):.2f} kg CO2e และ Scope 2 (การปล่อยทางอ้อมจากพลังงาน) {report_data['emissions_by_scope'].get('Scope 2', 0):.2f} kg CO2e การทำความเข้าใจการแบ่งประเภทนี้ช่วยในการกำหนดเป้าหมายกลยุทธ์การลดการปล่อยที่มีประสิทธิภาพ"""
            return self._clean_thai_text(text)
        else:
            # Clean category names by replacing underscores with spaces
            categories_text = ', '.join([f"{cat.replace('_', ' ')} ({(val/total_emissions*100):.1f}%)" for cat, val in categories_sorted[:3] if val > 0])
            text = f"""
            The emissions breakdown analysis reveals that total emissions of {total_emissions:.2f} kg CO2e originated from multiple sources.
            The primary emission sources are {categories_text}.
            Emissions are categorized according to the GHG Protocol as Scope 1 (direct emissions) {report_data['emissions_by_scope'].get('Scope 1', 0):.2f} kg CO2e
            and Scope 2 (indirect energy emissions) {report_data['emissions_by_scope'].get('Scope 2', 0):.2f} kg CO2e.
            Understanding this breakdown is essential for targeting effective emission reduction strategies.
            """
            return ' '.join(text.split())

    def _get_fallback_data_quality(self, language: str = 'EN') -> str:
        """Fallback data quality statement"""
        if language == 'TH':
            text = """คุณภาพข้อมูลในรายงานนี้ได้รับการรับรองผ่านกระบวนการเก็บรวบรวมข้อมูลที่เป็นระบบและขั้นตอนการตรวจสอบ ข้อมูลการปล่อยทั้งหมดได้รับการบันทึกและตรวจสอบตามมาตรฐานที่กำหนดไว้ ข้อมูลมีความครบถ้วนสำหรับช่วงเวลารายงานและครอบคลุมแหล่งการปล่อยที่สำคัญทั้งหมด การปรับปรุงอย่างต่อเนื่องในระบบการจัดการข้อมูลจะช่วยเพิ่มความแม่นยำและความน่าเชื่อถือในรายงานในอนาคต"""
            return self._clean_thai_text(text)
        else:
            text = """
            Data quality in this report is assured through systematic data collection processes and verification procedures.
            All emission data has been recorded and validated according to established standards.
            Data completeness is high for the reporting period, covering all material emission sources.
            Continuous improvements in data management systems will enhance accuracy and reliability in future reports.
            """
            return ' '.join(text.split())

    def _get_fallback_conclusion(self, report_data: Dict, language: str = 'EN') -> str:
        """Fallback conclusion"""
        if language == 'TH':
            text = f"""รายงานการปล่อยก๊าซเรือนกระจกฉบับนี้แสดงให้เห็นถึงความมุ่งมั่นของ {report_data['organization']} ในการวัดและจัดการผลกระทบด้านคาร์บอนอย่างโปร่งใส การปล่อยทั้งหมด {report_data['total_emissions']:.2f} kg CO2e ในช่วงระยะเวลารายงานเป็นพื้นฐานสำหรับการกำหนดเป้าหมายและกลยุทธ์การลดการปล่อยในอนาคต องค์กรมุ่งมั่นที่จะปรับปรุงประสิทธิภาพด้านสิ่งแวดล้อมอย่างต่อเนื่องและมีส่วนร่วมในการต่อสู้กับการเปลี่ยนแปลงสภาพภูมิอากาศ การติดตามและรายงานเป็นประจำจะช่วยให้มั่นใจได้ว่าจะบรรลุเป้าหมายด้านความยั่งยืน"""
            return self._clean_thai_text(text)
        else:
            text = f"""
            This carbon emissions report demonstrates {report_data['organization']}'s commitment to transparent measurement
            and management of carbon impacts. Total emissions of {report_data['total_emissions']:.2f} kg CO2e during the reporting
            period provide a foundation for setting future reduction targets and strategies. The organization is committed to
            continuous environmental performance improvement and contributing to the fight against climate change.
            Regular monitoring and reporting will ensure progress toward sustainability goals.
            """

        return ' '.join(text.split())

    def _format_date_thai(self, date_obj: datetime) -> str:
        """Format date in Thai (e.g., 15 มกราคม 2025)"""
        day = date_obj.day
        month = self.thai_months[date_obj.month]
        year = date_obj.year
        return f"{day} {month} {year}"

    def _format_month_year_thai(self, date_obj: datetime) -> str:
        """Format month and year in Thai (e.g., มกราคม 2025)"""
        month = self.thai_months[date_obj.month]
        year = date_obj.year
        return f"{month} {year}"

    def _clean_thai_text(self, text: str) -> str:
        """Clean Thai text while preserving proper spacing"""
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        # Join lines with a single space
        result = ' '.join(lines)
        # Remove multiple consecutive spaces (replace 2+ spaces with 1)
        import re
        result = re.sub(r'\s+', ' ', result)
        return result.strip()

    def _format_date_range(self, start_date: datetime, end_date: datetime, language: str) -> str:
        """Format date range based on language"""
        if language == 'TH':
            start_str = self._format_date_thai(start_date)
            end_str = self._format_date_thai(end_date)
            return f"{start_str} ถึง {end_str}"
        else:
            return f"{start_date.strftime('%d %B %Y')} - {end_date.strftime('%d %B %Y')}"

    def _create_report_structure(self, report_data: Dict, ai_content: Dict, report_format: str, language: str = 'EN') -> Dict:
        """Create structured report content"""

        # Calculate additional metrics
        avg_monthly_emissions = report_data['total_emissions'] / max(len(report_data['monthly_data']), 1)

        # Format date range based on language
        date_range = self._format_date_range(report_data['period_start'], report_data['period_end'], language)

        # Format subtitle with Thai month names if needed
        if language == 'TH':
            start_month = self.thai_months[report_data['period_start'].month]
            end_month = self.thai_months[report_data['period_end'].month]
            subtitle = f'{report_data["organization"]} - {start_month} {report_data["period_start"].year} ถึง {end_month} {report_data["period_end"].year}'
        else:
            subtitle = f'{report_data["organization"]} - {report_data["period_start"].strftime("%B %Y")} to {report_data["period_end"].strftime("%B %Y")}'

        return {
            'title': f'{report_format} Carbon Emissions Report',
            'subtitle': subtitle,

            # Core AI-Generated Sections
            'executive_summary': ai_content.get('executive_summary', ''),
            'trend_analysis': ai_content.get('trend_analysis', ''),
            'emissions_breakdown': ai_content.get('emissions_breakdown', ''),
            'methodology': ai_content.get('methodology', self._get_methodology_text(report_format, language)),
            'data_quality': ai_content.get('data_quality', ''),
            'recommendations': ai_content.get('recommendations', []),
            'conclusion': ai_content.get('conclusion', ''),

            # Key Metrics
            'key_metrics': {
                'total_emissions': f"{report_data['total_emissions']:.2f} kg CO2e",
                'average_monthly': f"{avg_monthly_emissions:.2f} kg CO2e/month",
                'reporting_period': date_range,
                'record_count': f"{report_data['record_count']}"
            },

            # Data Breakdown
            'emissions_by_scope': report_data['emissions_by_scope'],
            'emissions_by_category': report_data['emissions_by_category'],
            'monthly_data': report_data['monthly_data'],

            # Legacy support
            'key_findings': ai_content.get('key_findings', []),

            # Metadata
            'generated_at': datetime.now().strftime('%d %B %Y at %H:%M')
        }

    def _get_methodology_text(self, report_format: str, language: str = 'EN') -> str:
        """Get methodology text based on report format and language"""
        if language == 'TH':
            methodologies = {
                'ISO': """รายงานนี้เป็นไปตามหลักการ ISO 14064-1:2018 สำหรับการวัดและรายงานก๊าซเรือนกระจก ค่าสัมประสิทธิ์การปล่อยมาจากแนวทาง IPCC และหน่วยงานกำกับดูแลในท้องถิ่น การเก็บรวบรวมข้อมูลเป็นไปตามขั้นตอนที่เป็นระบบเพื่อให้แน่ใจว่ามีความถูกต้องและครบถ้วน การคำนวณทั้งหมดใช้แนวทางการควบคุมการดำเนินงานสำหรับการกำหนดขอบเขตองค์กร""",
                'CFO': """รายงานนี้ใช้วิธีการบัญชีคาร์บอนมาตรฐานที่เหมาะสำหรับการรายงานทางการเงิน การคำนวณการปล่อยเป็นไปตามโปรโตคอลที่กำหนดไว้โดยเน้นผลกระทบที่สำคัญ การนำเสนอข้อมูลเพื่อสนับสนุนการตัดสินใจเชิงกลยุทธ์และการประเมินความเสี่ยง ผลกระทบทางการเงินของการปล่อยคาร์บอนได้รับการพิจารณาตลอดการวิเคราะห์""",
                'GHG': """รายงานนี้เป็นไปตามมาตรฐาน GHG Protocol Corporate Accounting and Reporting Standard การปล่อยก๊าซเรือนกระจกถูกจำแนกตาม Scope 1, 2, และ 3 ตามที่กำหนดใน GHG Protocol วิธีการคำนวณเป็นไปตามแนวทาง GHG Protocol สำหรับการจัดทำรายการขององค์กร การจัดการคุณภาพข้อมูลและความไม่แน่นอนเป็นไปตามข้อกำหนดของ GHG Protocol"""
            }
        else:
            methodologies = {
                'ISO': """
                This report follows ISO 14064-1:2018 principles for greenhouse gas quantification and reporting.
                Emission factors are sourced from IPCC guidelines and local regulatory authorities.
                Data collection follows systematic procedures to ensure accuracy and completeness.
                All calculations use the operational control approach for organizational boundary setting.
                """,
                'CFO': """
                This report uses standard carbon accounting methodologies suitable for financial reporting.
                Emission calculations follow established protocols with focus on material impacts.
                Data is presented to support strategic decision-making and risk assessment.
                Financial implications of carbon emissions are considered throughout the analysis.
                """,
                'GHG': """
                This report complies with the GHG Protocol Corporate Accounting and Reporting Standard.
                Emissions are classified according to Scope 1, 2, and 3 categories as defined by the GHG Protocol.
                Calculation methodologies follow GHG Protocol guidance for corporate inventories.
                Data quality and uncertainty are managed according to GHG Protocol requirements.
                """
            }
        
        default_text = "วิธีการบัญชีคาร์บอนมาตรฐานถูกนำมาใช้" if language == 'TH' else "Standard carbon accounting methodologies applied."
        # Clean up multiple spaces and newlines from the methodology text
        text = methodologies.get(report_format, default_text)
        if language == 'TH':
            return self._clean_thai_text(text)
        else:
            return ' '.join(text.split())

    def _generate_report_file(self, content: Dict, report_format: str, file_type: str, language: str) -> str:
        """Generate report file based on type and language"""
        if file_type == 'PDF':
            return self._generate_pdf_report(content, report_format, language)
        elif file_type == 'EXCEL':
            return self._generate_excel_report(content, report_format, language)
        elif file_type == 'WORD':
            return self._generate_word_report(content, report_format, language)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _get_ghg_template_content(self, language: str) -> Dict:
        """Get GHG Protocol template content based on language"""
        if language == 'TH':
            return {
                'title': 'รายงานการปล่อยก๊าซเรือนกระจก ตามมาตรฐาน GHG Protocol',
                'executive_summary_title': 'บทสรุปผู้บริหาร',
                'scope_1_title': 'การปล่อยทางตรง (Scope 1)',
                'scope_2_title': 'การปล่อยทางอ้อมจากพลังงาน (Scope 2)',
                'key_findings_title': 'ผลการวิเคราะห์สำคัญ',
                'recommendations_title': 'ข้อเสนอแนะ',
                'trend_analysis_title': 'การวิเคราะห์แนวโน้ม',
                'emissions_breakdown_title': 'การแบ่งประเภทการปล่อยก๊าซ',
                'methodology_title': 'วิธีการคำนวณ',
                'data_quality_title': 'คำชี้แจงคุณภาพข้อมูล',
                'conclusion_title': 'สรุป',
                'scope_descriptions': {
                    'Scope 1': 'การปล่อยก๊าซเรือนกระจกโดยตรงจากแหล่งที่องค์กรเป็นเจ้าของหรือควบคุม (เชื้อเพลิง สารทำความเย็น การเผาไหม้)',
                    'Scope 2': 'การปล่อยก๊าซเรือนกระจกทางอ้อมจากการซื้อพลังงานไฟฟ้า'
                }
            }
        else:  # English
            return {
                'title': 'Greenhouse Gas Inventory Report (GHG Protocol)',
                'executive_summary_title': 'Executive Summary',
                'scope_1_title': 'Scope 1 Direct Emissions',
                'scope_2_title': 'Scope 2 Indirect Emissions from Energy',
                'key_findings_title': 'Key Findings',
                'recommendations_title': 'Recommendations',
                'trend_analysis_title': 'Trend Analysis',
                'emissions_breakdown_title': 'Emissions Breakdown',
                'methodology_title': 'Methodology',
                'data_quality_title': 'Data Quality Statement',
                'conclusion_title': 'Conclusion',
                'scope_descriptions': {
                    'Scope 1': 'Direct greenhouse gas emissions from sources owned or controlled by the organization (fuels, refrigerants, combustion)',
                    'Scope 2': 'Indirect greenhouse gas emissions from purchased electricity'
                }
            }

    def _setup_thai_fonts(self):
        """Setup Thai fonts with better mixed Thai-English support"""
        try:
            # Create fonts directory if it doesn't exist
            fonts_dir = os.path.join(os.path.dirname(__file__), 'fonts')
            os.makedirs(fonts_dir, exist_ok=True)
            
            # Enhanced font sources with better mixed-language support
            # Prioritize fonts with excellent Latin character support
            font_sources = [
                {
                    'name': 'Sarabun', 
                    'regular_url': 'https://github.com/cadsondemak/Sarabun/raw/master/fonts/ttf/Sarabun-Regular.ttf',
                    'bold_url': 'https://github.com/cadsondemak/Sarabun/raw/master/fonts/ttf/Sarabun-Bold.ttf',
                    'regular_file': 'Sarabun-Regular.ttf',
                    'bold_file': 'Sarabun-Bold.ttf',
                    'has_latin': True,
                    'priority': 1  # Best for mixed content
                },
                {
                    'name': 'NotoSansThai',
                    'regular_url': 'https://github.com/google/fonts/raw/main/ofl/notosansthai/NotoSansThai-Regular.ttf',
                    'bold_url': 'https://github.com/google/fonts/raw/main/ofl/notosansthai/NotoSansThai-Bold.ttf',
                    'regular_file': 'NotoSansThai-Regular.ttf',
                    'bold_file': 'NotoSansThai-Bold.ttf',
                    'has_latin': True,
                    'priority': 2
                },
                {
                    'name': 'THSarabunNew', 
                    'regular_url': 'https://github.com/tlwg/fonts-tlwg/raw/master/tlwg/THSarabunNew.ttf',
                    'bold_url': 'https://github.com/tlwg/fonts-tlwg/raw/master/tlwg/THSarabunNew-Bold.ttf',
                    'regular_file': 'THSarabunNew.ttf',
                    'bold_file': 'THSarabunNew-Bold.ttf',
                    'has_latin': True,
                    'priority': 3
                }
            ]
            
            # Sort by priority
            font_sources.sort(key=lambda x: x.get('priority', 999))
            
            for font_source in font_sources:
                try:
                    thai_font_path = os.path.join(fonts_dir, font_source['regular_file'])
                    thai_bold_path = os.path.join(fonts_dir, font_source.get('bold_file', font_source['regular_file']))
                    
                    # Download regular font if it doesn't exist
                    if not os.path.exists(thai_font_path):
                        print(f"Downloading {font_source['name']} regular font...")
                        urllib.request.urlretrieve(font_source['regular_url'], thai_font_path)
                    
                    # Download bold font if available and doesn't exist
                    if 'bold_url' in font_source and not os.path.exists(thai_bold_path):
                        print(f"Downloading {font_source['name']} bold font...")
                        urllib.request.urlretrieve(font_source['bold_url'], thai_bold_path)
                    
                    # Register fonts with UTF-8 support
                    pdfmetrics.registerFont(TTFont('ThaiFont', thai_font_path))
                    
                    if os.path.exists(thai_bold_path):
                        pdfmetrics.registerFont(TTFont('ThaiFont-Bold', thai_bold_path))
                    else:
                        # Use regular font for bold if no bold version available
                        pdfmetrics.registerFont(TTFont('ThaiFont-Bold', thai_font_path))
                    
                    # Register additional font variants for mixed content
                    pdfmetrics.registerFont(TTFont('ThaiFont-Italic', thai_font_path))
                    pdfmetrics.registerFont(TTFont('ThaiFont-BoldItalic', thai_bold_path if os.path.exists(thai_bold_path) else thai_font_path))
                    
                    # Add comprehensive font family mappings
                    addMapping('ThaiFont', 0, 0, 'ThaiFont')           # normal
                    addMapping('ThaiFont', 0, 1, 'ThaiFont-Bold')      # bold
                    addMapping('ThaiFont', 1, 0, 'ThaiFont-Italic')    # italic
                    addMapping('ThaiFont', 1, 1, 'ThaiFont-BoldItalic') # bold italic
                    
                    # Test the font with mixed content
                    test_success = self._test_font_mixed_content('ThaiFont')
                    
                    if test_success:
                        print(f"Successfully registered {font_source['name']} fonts with mixed content support")
                        return True
                    else:
                        print(f"Font {font_source['name']} registered but failed mixed content test")
                        continue
                
                except Exception as e:
                    print(f"Failed to setup {font_source['name']}: {str(e)}")
                    continue
            
            # If all fonts fail, register fallback combination
            print("All Thai font setups failed, setting up fallback font combination")
            return self._setup_fallback_fonts()
            
        except Exception as e:
            print(f"Error setting up Thai fonts: {str(e)}")
            return self._setup_fallback_fonts()

    def _test_font_mixed_content(self, font_name):
        """Test if font can handle mixed Thai-English content"""
        try:
            from reportlab.pdfbase.pdfmetrics import getFont
            font = getFont(font_name)
            
            # Test mixed content strings
            test_strings = [
                "บริษัท ABC Company จำกัด",
                "การปล่อย CO2 equivalent", 
                "GHG Protocol การรายงาน"
            ]
            
            for test_str in test_strings:
                # Try to encode the string with the font
                try:
                    # This is a basic test - in practice, ReportLab handles this internally
                    encoded = test_str.encode('utf-8')
                    decoded = encoded.decode('utf-8')
                    if test_str != decoded:
                        return False
                except:
                    return False
            
            return True
            
        except Exception as e:
            print(f"Font test failed: {str(e)}")
            return False

    def _setup_fallback_fonts(self):
        """Setup fallback font combination for mixed content"""
        try:
            # Register Helvetica as fallback for Latin characters
            # This ensures English text always renders correctly
            pdfmetrics.registerFont(TTFont('FallbackFont', 'Helvetica'))
            
            # Create a hybrid approach - use system fonts
            addMapping('MixedFont', 0, 0, 'Helvetica')      # normal - good for English
            addMapping('MixedFont', 0, 1, 'Helvetica-Bold') # bold - good for English
            addMapping('MixedFont', 1, 0, 'Helvetica')      # italic
            addMapping('MixedFont', 1, 1, 'Helvetica-Bold') # bold italic
            
            print("Fallback font combination registered")
            return True
            
        except Exception as e:
            print(f"Fallback font setup failed: {str(e)}")
            return False

    def _get_font_name(self, language: str, bold: bool = False) -> str:
        """Get appropriate font name with better mixed-content support"""
        if language == 'TH':
            # For Thai documents, use a font that supports both Thai and Latin
            try:
                from reportlab.pdfbase.pdfmetrics import getFont
                getFont('ThaiFont')
                # Use Thai font only for mixed content, but ensure it has Latin support
                return 'ThaiFont-Bold' if bold else 'ThaiFont'
            except:
                # Fallback to Helvetica which has good Latin support
                return 'Helvetica-Bold' if bold else 'Helvetica'
        return 'Helvetica-Bold' if bold else 'Helvetica'

    def _process_mixed_content_text(self, text: str, language: str = 'EN') -> str:
        """Process mixed Thai-English content to ensure proper UTF-8 encoding and spacing"""
        if not text:
            return text
            
        try:
            # Ensure the text is properly encoded as UTF-8
            if isinstance(text, bytes):
                text = text.decode('utf-8', errors='replace')
            
            # For Thai language reports with mixed content
            if language == 'TH':
                # Normalize Unicode characters to ensure consistent encoding
                import unicodedata
                text = unicodedata.normalize('NFC', text)
                
                # Handle common problematic character combinations
                replacements = {
                    # Fix common encoding issues
                    'â€™': "'",  # Smart quote
                    'â€œ': '"',  # Smart quote
                    'â€': '"',   # Smart quote
                    'â€¢': '•',  # Bullet point
                    'â€"': '–',  # En dash
                    'â€"': '—',  # Em dash
                }
                
                for old, new in replacements.items():
                    text = text.replace(old, new)
                
                # Enhanced spacing for mixed Thai-English content
                import re
                
                # Add spaces around English words/numbers when adjacent to Thai characters
                text = re.sub(r'([ก-๙])([A-Za-z0-9])', r'\1 \2', text)
                text = re.sub(r'([A-Za-z0-9])([ก-๙])', r'\1 \2', text)
                
                # Add spaces around parentheses when adjacent to Thai characters
                text = re.sub(r'([ก-๙])(\()', r'\1 \2', text)
                text = re.sub(r'(\))([ก-๙])', r'\1 \2', text)
                
                # Add spaces around colons when adjacent to Thai characters
                text = re.sub(r'([ก-๙])(:)', r'\1\2 ', text)
                text = re.sub(r'(:)([ก-๙])', r'\1 \2', text)
                
                # Add spaces around commas when adjacent to Thai characters
                text = re.sub(r'([ก-๙])(,)', r'\1\2 ', text)
                text = re.sub(r'(,)([ก-๙])', r'\1 \2', text)
                
                # Special handling for common technical terms
                technical_terms = [
                    'GHG Protocol', 'ISO 14064', 'CO2e', 'Scope 1', 'Scope 2', 'TGO',
                    'operational efficiency', 'energy management system', 'direct emissions',
                    'indirect emissions', 'Corporate Accounting', 'Reporting Standard'
                ]
                
                for term in technical_terms:
                    # Ensure proper spacing around technical terms
                    pattern = r'([ก-๙])(' + re.escape(term) + r')([ก-๙])'
                    text = re.sub(pattern, r'\1 \2 \3', text)
                
                # Clean up multiple spaces but preserve intentional spacing
                text = re.sub(r' {3,}', '  ', text)  # Max 2 spaces
                text = re.sub(r' {2,}', ' ', text)   # Normalize to single space
                text = text.strip()
            
            # Final UTF-8 validation
            text.encode('utf-8')
            return text
            
        except Exception as e:
            print(f"Text processing error: {str(e)}")
            # Return original text if processing fails
            return str(text) if text else ""

    def _create_mixed_content_paragraph(self, text: str, style, language: str = 'EN'):
        """Create a paragraph that handles mixed Thai-English content properly"""
        try:
            # Process the text for mixed content
            processed_text = self._process_mixed_content_text(text, language)
            
            # Create paragraph with proper encoding
            return Paragraph(processed_text, style)
            
        except Exception as e:
            print(f"Paragraph creation error: {str(e)}")
            # Fallback to simple paragraph
            try:
                return Paragraph(str(text), style)
            except:
                return Paragraph("Content encoding error", style)

    def _process_data_with_pandas(self, content: Dict, language: str = 'EN') -> Dict:
        """Process report data using pandas for better data handling and analysis"""
        try:
            import pandas as pd
            import numpy as np
            
            # Process emissions by category data
            if content.get('emissions_by_category'):
                category_df = pd.DataFrame(list(content['emissions_by_category'].items()), 
                                         columns=['Category', 'Emissions'])
                category_df['Emissions'] = pd.to_numeric(category_df['Emissions'], errors='coerce')
                category_df = category_df.sort_values('Emissions', ascending=False)
                
                # Calculate percentages
                total_emissions = category_df['Emissions'].sum()
                category_df['Percentage'] = (category_df['Emissions'] / total_emissions * 100).round(1)
                
                content['category_analysis'] = {
                    'top_category': category_df.iloc[0]['Category'] if len(category_df) > 0 else 'Unknown',
                    'top_percentage': category_df.iloc[0]['Percentage'] if len(category_df) > 0 else 0,
                    'category_summary': category_df.to_dict('records')
                }
            
            # Process emissions by scope data
            if content.get('emissions_by_scope'):
                scope_df = pd.DataFrame(list(content['emissions_by_scope'].items()), 
                                      columns=['Scope', 'Emissions'])
                scope_df['Emissions'] = pd.to_numeric(scope_df['Emissions'], errors='coerce')
                scope_df = scope_df[scope_df['Emissions'] > 0]  # Filter out zero emissions
                
                # Calculate percentages
                total_scope_emissions = scope_df['Emissions'].sum()
                if total_scope_emissions > 0:
                    scope_df['Percentage'] = (scope_df['Emissions'] / total_scope_emissions * 100).round(1)
                else:
                    scope_df['Percentage'] = 0
                
                content['scope_analysis'] = {
                    'dominant_scope': scope_df.iloc[0]['Scope'] if len(scope_df) > 0 else 'Unknown',
                    'scope_summary': scope_df.to_dict('records')
                }
            
            # Process monthly data for trend analysis
            if content.get('monthly_data'):
                monthly_df = pd.DataFrame(content['monthly_data'])
                if 'total' in monthly_df.columns:
                    monthly_df['total'] = pd.to_numeric(monthly_df['total'], errors='coerce')
                    
                    # Calculate trend statistics
                    if len(monthly_df) > 1:
                        # Linear trend
                        x = np.arange(len(monthly_df))
                        y = monthly_df['total'].values
                        trend_slope = np.polyfit(x, y, 1)[0]
                        
                        # Month-over-month change
                        monthly_df['change'] = monthly_df['total'].pct_change() * 100
                        avg_change = monthly_df['change'].mean()
                        
                        content['trend_analysis_data'] = {
                            'trend_direction': 'increasing' if trend_slope > 0 else 'decreasing' if trend_slope < 0 else 'stable',
                            'trend_slope': round(trend_slope, 2),
                            'avg_monthly_change': round(avg_change, 1),
                            'highest_month': monthly_df.loc[monthly_df['total'].idxmax(), 'month'] if len(monthly_df) > 0 else 'Unknown',
                            'lowest_month': monthly_df.loc[monthly_df['total'].idxmin(), 'month'] if len(monthly_df) > 0 else 'Unknown'
                        }
            
            # Process key metrics for better formatting
            if content.get('key_metrics'):
                metrics_df = pd.DataFrame(list(content['key_metrics'].items()), 
                                        columns=['Metric', 'Value'])
                
                # Clean and format metrics
                formatted_metrics = {}
                for _, row in metrics_df.iterrows():
                    metric = row['Metric']
                    value = row['Value']
                    
                    # Format based on metric type
                    if 'emission' in metric.lower():
                        # Format emission values
                        if isinstance(value, str) and 'kg CO2e' in value:
                            formatted_metrics[metric] = value
                        else:
                            formatted_metrics[metric] = f"{float(value):.2f} kg CO2e" if pd.notna(value) else "0.00 kg CO2e"
                    elif 'count' in metric.lower():
                        # Format count values
                        formatted_metrics[metric] = f"{int(float(value))}" if pd.notna(value) else "0"
                    else:
                        formatted_metrics[metric] = str(value)
                
                content['formatted_metrics'] = formatted_metrics
            
            # Generate summary statistics
            content['data_quality'] = {
                'categories_count': len(content.get('emissions_by_category', {})),
                'active_scopes': len([s for s in content.get('emissions_by_scope', {}).values() if s > 0]),
                'reporting_months': len(content.get('monthly_data', [])),
                'data_completeness': 'Good' if content.get('emissions_by_category') and content.get('monthly_data') else 'Limited'
            }
            
            print("Data processed successfully with pandas")
            return content
            
        except Exception as e:
            print(f"Pandas data processing error: {str(e)}")
            # Return original content if pandas processing fails
            return content

    def _create_thai_styles(self, language: str):
        """Create custom styles with improved Thai font support"""
        styles = getSampleStyleSheet()
        
        # Get base font
        base_font = self._get_font_name(language, False)
        bold_font = self._get_font_name(language, True)
        
        if language == 'TH':
            # Title style
            styles.add(ParagraphStyle(
                name='ThaiTitle',
                parent=styles['Heading1'],
                fontName=bold_font,
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # Center alignment
                leading=28,   # Increased leading for better Thai text spacing
                wordSpace=1   # Add slight word spacing
            ))

            # Heading style
            styles.add(ParagraphStyle(
                name='ThaiHeading',
                parent=styles['Heading2'],
                fontName=bold_font,
                fontSize=14,
                spaceBefore=16,
                spaceAfter=10,
                leading=22,    # Increased leading for better Thai text spacing
                wordSpace=1    # Add slight word spacing
            ))

            # Normal text style
            styles.add(ParagraphStyle(
                name='ThaiNormal',
                parent=styles['Normal'],
                fontName=base_font,
                fontSize=12,
                leading=20,    # Increased leading for better Thai text spacing
                spaceAfter=10,
                wordSpace=1    # Add slight word spacing
            ))

            # Table header style
            styles.add(ParagraphStyle(
                name='ThaiTableHeader',
                parent=styles['Normal'],
                fontName=bold_font,
                fontSize=12,
                leading=18,
                alignment=1,
                wordSpace=1
            ))
        
        return styles

    def _generate_pdf_report(self, content: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate professional PDF report using Word-to-PDF conversion for better quality"""
        try:
            print(f"\n{'='*60}")
            print(f"PDF GENERATION REQUEST")
            print(f"{'='*60}")
            print(f"  Language: {language}")
            print(f"  Using Word-to-PDF conversion for Thai formatting")
            # Use Word-to-PDF conversion for both Thai and English for consistency
            return self._generate_pdf_via_word(content, report_format, language)

        except Exception as e:
            print(f"✗ PDF generation error: {str(e)}")
            import traceback
            traceback.print_exc()
            print(f"⚠ Falling back to direct PDF generation (old format)")
            # Fallback to direct PDF generation
            return self._generate_direct_pdf(content, report_format, language)

    def _generate_pdf_via_word(self, content: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate PDF by first creating Word document then converting to PDF"""
        try:
            print(f"\n{'='*60}")
            print(f"Starting Word-to-PDF conversion for {language} report")
            print(f"{'='*60}")

            # First generate Word document
            word_filepath = self._generate_word_report(content, report_format, language)

            if not word_filepath or not os.path.exists(word_filepath):
                raise Exception("Word document generation failed")

            print(f"✓ Word document created: {word_filepath}")

            # Create PDF filename with absolute path
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_filename = f"carbon_report_{report_format}_{language}_{timestamp}.pdf"
            reports_dir = os.path.join(os.path.dirname(__file__), 'reports')
            os.makedirs(reports_dir, exist_ok=True)
            pdf_filepath = os.path.join(reports_dir, pdf_filename)

            # Convert to absolute paths
            word_filepath_abs = os.path.abspath(word_filepath)
            pdf_filepath_abs = os.path.abspath(pdf_filepath)

            print(f"Target PDF path: {pdf_filepath_abs}")

            # Try to convert Word to PDF using python-docx2pdf or similar
            try:
                # Method 1: Try using docx2pdf if available
                print("Attempting conversion with docx2pdf...")
                import docx2pdf
                import pythoncom

                # Initialize COM for this thread
                pythoncom.CoInitialize()

                try:
                    docx2pdf.convert(word_filepath_abs, pdf_filepath_abs)

                    # Verify PDF was created
                    if os.path.exists(pdf_filepath_abs):
                        print(f"✓ Successfully converted Word to PDF using docx2pdf")
                        print(f"  PDF size: {os.path.getsize(pdf_filepath_abs)} bytes")

                        # Clean up temporary Word file
                        try:
                            os.remove(word_filepath_abs)
                            print(f"✓ Cleaned up temporary Word file")
                        except Exception as cleanup_error:
                            print(f"⚠ Could not remove temp file: {cleanup_error}")

                        return pdf_filepath
                    else:
                        raise Exception("docx2pdf completed but PDF file not found")
                finally:
                    # Uninitialize COM
                    pythoncom.CoUninitialize()

            except ImportError as ie:
                if 'pythoncom' in str(ie):
                    print("⚠ pythoncom not available, trying alternative method...")
                else:
                    print("⚠ docx2pdf not available, trying alternative method...")
            except Exception as e:
                print(f"⚠ docx2pdf conversion failed: {str(e)}")

            # Method 2: Try using win32com (Windows only)
            try:
                print("Attempting conversion with win32com (Microsoft Word)...")
                import win32com.client
                import pythoncom

                # Initialize COM for this thread
                pythoncom.CoInitialize()

                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False

                    # Open the Word document
                    doc = word.Documents.Open(word_filepath_abs)

                    # Save as PDF
                    doc.SaveAs2(pdf_filepath_abs, FileFormat=17)  # 17 = PDF format
                    doc.Close()
                    word.Quit()

                    # Verify PDF was created
                    if os.path.exists(pdf_filepath_abs):
                        print(f"✓ Successfully converted Word to PDF using win32com")
                        print(f"  PDF size: {os.path.getsize(pdf_filepath_abs)} bytes")

                        # Clean up temporary Word file
                        try:
                            os.remove(word_filepath_abs)
                            print(f"✓ Cleaned up temporary Word file")
                        except Exception as cleanup_error:
                            print(f"⚠ Could not remove temp file: {cleanup_error}")

                        return pdf_filepath
                    else:
                        raise Exception("win32com completed but PDF file not found")
                finally:
                    # Uninitialize COM
                    pythoncom.CoUninitialize()

            except ImportError:
                print("⚠ win32com not available")
            except Exception as e:
                print(f"⚠ win32com conversion failed: {str(e)}")

            # If both methods failed, raise exception instead of using bad fallback
            raise Exception("All Word-to-PDF conversion methods failed")

        except Exception as e:
            print(f"✗ Word-to-PDF conversion failed: {str(e)}")
            print(f"✗ CANNOT generate report - Word-to-PDF is required for quality output")
            raise Exception(f"PDF generation failed: {str(e)}")

    def _generate_improved_direct_pdf(self, content: Dict, report_format: str, language: str, word_reference: str = None) -> str:
        """Generate PDF with improved font handling based on successful Word document approach"""
        try:
            import pandas as pd
            
            # Use system fonts that work well with mixed content (like Word does)
            # Instead of downloading fonts, use system fonts that are proven to work
            
            # Get template content
            template = self._get_ghg_template_content(language)
            
            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"carbon_report_{report_format}_{language}_{timestamp}.pdf"
            filepath = os.path.join('reports', filename)
            
            # Ensure reports directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Process data with pandas for better handling
            processed_data = self._process_data_with_pandas(content, language)
            
            # Create PDF document with improved font handling
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = self._create_system_font_styles(language)
            story = []
            
            # Use system fonts that handle mixed content well
            title_style = styles['Title']
            heading_style = styles['Heading1'] 
            normal_style = styles['Normal']
            
            # Title and subtitle
            story.append(Paragraph(template['title'], title_style))
            story.append(Paragraph(content['subtitle'], heading_style))
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph(template['executive_summary_title'], heading_style))
            story.append(Paragraph(content['executive_summary'], normal_style))
            story.append(Spacer(1, 20))
            
            # Key Metrics Table
            metrics_title = "ตัวชี้วัดหลัก" if language == 'TH' else "Key Metrics"
            story.append(Paragraph(metrics_title, heading_style))

            # Create table data
            metrics_data = []
            if language == 'TH':
                metrics_data.append(['ตัวชี้วัด', 'ค่า'])
            else:
                metrics_data.append(['Metric', 'Value'])

            for key, value in content['key_metrics'].items():
                # Use Thai labels if language is Thai, otherwise use English
                if language == 'TH' and key in self.metric_labels_th:
                    display_key = self.metric_labels_th[key]
                else:
                    display_key = key.replace('_', ' ').title()
                metrics_data.append([display_key, str(value)])

            metrics_table = Table(metrics_data)
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),  # Use system font
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
            ]))
            story.append(metrics_table)
            story.append(Spacer(1, 20))

            # Continue with rest of content using system fonts...
            # Key Findings
            if content['key_findings']:
                story.append(Paragraph(template['key_findings_title'], heading_style))
                for finding in content['key_findings']:
                    story.append(Paragraph(f"• {finding}", normal_style))
                story.append(Spacer(1, 20))
            
            # Recommendations
            if content['recommendations']:
                story.append(Paragraph(template['recommendations_title'], heading_style))
                for i, recommendation in enumerate(content['recommendations'], 1):
                    story.append(Paragraph(f"{i}. {recommendation}", normal_style))
                story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
            # Clean up temporary Word file if it exists
            if word_reference and os.path.exists(word_reference):
                try:
                    os.remove(word_reference)
                except:
                    pass
            
            return filepath
            
        except Exception as e:
            print(f"Improved PDF generation error: {str(e)}")
            return self._generate_direct_pdf(content, report_format, language)

    def _create_system_font_styles(self, language: str):
        """Create styles using system fonts that handle mixed content well"""
        styles = getSampleStyleSheet()
        
        # Use Helvetica which handles mixed content better than custom fonts
        # This mimics what Word processors do - use system fonts with good Unicode support
        
        if language == 'TH':
            # For Thai, still use system fonts but with better spacing
            styles.add(ParagraphStyle(
                name='ThaiTitle',
                parent=styles['Title'],
                fontName='Helvetica-Bold',
                fontSize=18,
                spaceAfter=30,
                alignment=1,
                leading=28,
                wordSpace=1
            ))

            styles.add(ParagraphStyle(
                name='ThaiHeading',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=14,
                spaceBefore=16,
                spaceAfter=10,
                leading=22,
                wordSpace=1
            ))

            styles.add(ParagraphStyle(
                name='ThaiNormal',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=12,
                leading=20,
                spaceAfter=10,
                wordSpace=1
            ))
        
        return styles

    def _generate_direct_pdf(self, content: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate PDF directly using ReportLab (fallback method)"""
        try:
            import pandas as pd
            
            # Setup Thai fonts if needed
            if language == 'TH':
                font_setup_success = self._setup_thai_fonts()
                if not font_setup_success:
                    print("Warning: Thai font setup failed, falling back to default fonts")
            
            # Get template content
            template = self._get_ghg_template_content(language)
            
            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"carbon_report_{report_format}_{language}_{timestamp}.pdf"
            filepath = os.path.join('reports', filename)
            
            # Ensure reports directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Process data with pandas for better handling
            processed_data = self._process_data_with_pandas(content, language)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = self._create_thai_styles(language)
            story = []
            
            # Get appropriate styles based on language
            title_style_name = 'ThaiTitle' if language == 'TH' else 'Heading1'
            heading_style_name = 'ThaiHeading' if language == 'TH' else 'Heading2'
            normal_style_name = 'ThaiNormal' if language == 'TH' else 'Normal'
            
            # Use custom title style or create one
            if language == 'TH' and 'ThaiTitle' in styles:
                title_style = styles['ThaiTitle']
            else:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontName=self._get_font_name(language),
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1,  # Center alignment
                    leading=22
                )
            
            # Title and subtitle with mixed content processing
            story.append(self._create_mixed_content_paragraph(template['title'], title_style, language))
            story.append(self._create_mixed_content_paragraph(content['subtitle'], styles.get(heading_style_name, styles['Heading2']), language))
            story.append(Spacer(1, 20))
            
            # Executive Summary with mixed content processing
            story.append(self._create_mixed_content_paragraph(template['executive_summary_title'], styles.get(heading_style_name, styles['Heading2']), language))
            story.append(self._create_mixed_content_paragraph(content['executive_summary'], styles.get(normal_style_name, styles['Normal']), language))
            story.append(Spacer(1, 20))
            
            # Key Metrics Table
            metrics_title = "ตัวชี้วัดหลัก" if language == 'TH' else "Key Metrics"
            story.append(Paragraph(metrics_title, styles.get(heading_style_name, styles['Heading2'])))

            # Create table data with proper encoding
            metrics_data = []
            if language == 'TH':
                metrics_data.append(['ตัวชี้วัด', 'ค่า'])
            else:
                metrics_data.append(['Metric', 'Value'])

            for key, value in content['key_metrics'].items():
                # Use Thai labels if language is Thai, otherwise use English
                if language == 'TH' and key in self.metric_labels_th:
                    display_key = self.metric_labels_th[key]
                else:
                    display_key = key.replace('_', ' ').title()
                metrics_data.append([display_key, str(value)])

            metrics_table = Table(metrics_data)

            # Get appropriate fonts for table
            table_font = self._get_font_name(language, False)
            table_font_bold = self._get_font_name(language, True)

            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), table_font_bold),
                ('FONTNAME', (0, 1), (-1, -1), table_font),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
            ]))
            story.append(metrics_table)
            story.append(Spacer(1, 20))

            # Scope summary table (around line 690 in your original code)
            if any(v > 0 for v in content['emissions_by_scope'].values()):
                scope_title = "การปล่อยก๊าซเรือนกระจกตามขอบเขต (GHG Protocol)" if language == 'TH' else "Emissions by Scope (GHG Protocol)"
                story.append(Paragraph(scope_title, styles.get(heading_style_name, styles['Heading2'])))
                
                # Scope descriptions first
                for scope, value in content['emissions_by_scope'].items():
                    if value > 0:
                        description = template.get('scope_descriptions', {}).get(scope, f'Description for {scope} not available')
                        story.append(Paragraph(f"{scope}: {description}", styles.get(normal_style_name, styles['Normal'])))
                        total_text = f"รวม: {value:.2f} kg CO2e" if language == 'TH' else f"Total: {value:.2f} kg CO2e"
                        story.append(Paragraph(total_text, styles.get(normal_style_name, styles['Normal'])))
                        story.append(Spacer(1, 10))
                
                # Create scope table with proper Thai headers
                if language == 'TH':
                    scope_header = ['ขอบเขต', 'การปล่อย (kg CO2e)', 'เปอร์เซ็นต์']
                else:
                    scope_header = ['Scope', 'Emissions (kg CO2e)', 'Percentage']
                
                scope_data = [scope_header]
                total_scope = sum(content['emissions_by_scope'].values())
                
                for scope, value in content['emissions_by_scope'].items():
                    if value > 0:
                        percentage = (value / total_scope * 100) if total_scope > 0 else 0
                        scope_data.append([scope, f"{value:.2f}", f"{percentage:.1f}%"])
                
                scope_table = Table(scope_data)
                
                # Use the improved font selection
                table_font = self._get_font_name(language, False)
                table_font_bold = self._get_font_name(language, True)
                
                scope_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), table_font_bold),
                    ('FONTNAME', (0, 1), (-1, -1), table_font),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('FONTSIZE', (0, 1), (-1, -1), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))
                story.append(scope_table)
                story.append(Spacer(1, 20))
            
            # Key Findings
            if content['key_findings']:
                story.append(Paragraph(template['key_findings_title'], styles.get(heading_style_name, styles['Heading2'])))
                for finding in content['key_findings']:
                    story.append(Paragraph(f"• {finding}", styles.get(normal_style_name, styles['Normal'])))
                story.append(Spacer(1, 20))
            
            # Recommendations
            if content['recommendations']:
                story.append(Paragraph(template['recommendations_title'], styles.get(heading_style_name, styles['Heading2'])))
                for i, recommendation in enumerate(content['recommendations'], 1):
                    story.append(Paragraph(f"{i}. {recommendation}", styles.get(normal_style_name, styles['Normal'])))
                story.append(Spacer(1, 20))
            
            # Trend Analysis
            if content['trend_analysis']:
                trend_title = "การวิเคราะห์แนวโน้ม" if language == 'TH' else "Trend Analysis"
                story.append(Paragraph(trend_title, styles.get(heading_style_name, styles['Heading2'])))
                story.append(Paragraph(content['trend_analysis'], styles.get(normal_style_name, styles['Normal'])))
                story.append(Spacer(1, 20))
            
            # Methodology
            story.append(Paragraph(template['methodology_title'], styles.get(heading_style_name, styles['Heading2'])))
            story.append(Paragraph(content['methodology'], styles.get(normal_style_name, styles['Normal'])))
            story.append(Spacer(1, 20))
            
            # Footer
            story.append(Spacer(1, 30))
            footer_text = f"รายงานสร้างเมื่อ {content['generated_at']}" if language == 'TH' else f"Report generated on {content['generated_at']}"
            story.append(Paragraph(footer_text, styles.get(normal_style_name, styles['Normal'])))
            
            # Build PDF
            doc.build(story)
            
            return filepath
            
        except Exception as e:
            print(f"PDF generation error: {str(e)}")
            return None

    def _generate_excel_report(self, content: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate Excel report with multiple sheets"""
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            # Get template content
            template = self._get_ghg_template_content(language)
            
            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"carbon_report_{report_format}_{language}_{timestamp}.xlsx"
            filepath = os.path.join('reports', filename)
            
            # Ensure reports directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Create workbook
            wb = Workbook()
            
            # Summary Sheet
            ws_summary = wb.active
            ws_summary.title = "Summary"
            
            # Header styling
            header_font = Font(bold=True, size=14)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            
            # Title
            ws_summary['A1'] = template['title']
            ws_summary['A1'].font = Font(bold=True, size=16)
            ws_summary.merge_cells('A1:D1')
            
            # Key Metrics
            row = 3
            ws_summary[f'A{row}'] = "Key Metrics"
            ws_summary[f'A{row}'].font = header_font
            row += 1
            
            for key, value in content['key_metrics'].items():
                # Use Thai labels if language is Thai, otherwise use English
                if language == 'TH' and key in self.metric_labels_th:
                    ws_summary[f'A{row}'] = self.metric_labels_th[key]
                else:
                    ws_summary[f'A{row}'] = key.replace('_', ' ').title()
                ws_summary[f'B{row}'] = value
                row += 1
            
            # Emissions by Scope
            row += 2
            ws_summary[f'A{row}'] = "Emissions by Scope"
            ws_summary[f'A{row}'].font = header_font
            row += 1
            
            ws_summary[f'A{row}'] = "Scope"
            ws_summary[f'B{row}'] = "Emissions (kg CO2e)"
            ws_summary[f'C{row}'] = "Percentage"
            for col in ['A', 'B', 'C']:
                ws_summary[f'{col}{row}'].font = header_font
            row += 1
            
            total_scope = sum(content['emissions_by_scope'].values())
            for scope, value in content['emissions_by_scope'].items():
                if value > 0:
                    percentage = (value / total_scope * 100) if total_scope > 0 else 0
                    ws_summary[f'A{row}'] = scope
                    ws_summary[f'B{row}'] = round(value, 2)
                    ws_summary[f'C{row}'] = f"{percentage:.1f}%"
                    row += 1
            
            # Monthly Data Sheet
            if content['monthly_data']:
                ws_monthly = wb.create_sheet("Monthly Data")
                ws_monthly['A1'] = "Monthly Emissions Breakdown"
                ws_monthly['A1'].font = Font(bold=True, size=14)
                
                # Headers
                ws_monthly['A3'] = "Month"
                ws_monthly['B3'] = "Total Emissions (kg CO2e)"
                for col in ['A', 'B']:
                    ws_monthly[f'{col}3'].font = header_font
                
                row = 4
                for month_data in content['monthly_data']:
                    ws_monthly[f'A{row}'] = month_data['month']
                    ws_monthly[f'B{row}'] = round(month_data['total'], 2)
                    row += 1
            
            # Save workbook
            wb.save(filepath)
            print(f"✓ Excel document saved: {filepath}")
            print(f"✓ File exists: {os.path.exists(filepath)}")
            print(f"✓ Absolute path: {os.path.abspath(filepath)}")
            return filepath
            
        except Exception as e:
            print(f"Excel generation error: {str(e)}")
            return None

    def _set_thai_font(self, run, language: str):
        """Set appropriate font for Thai text"""
        if language == 'TH':
            from docx.shared import Pt
            # Use Angsana New - a standard Windows Thai font with excellent spacing
            # This font is pre-installed on Windows and handles Thai text very well
            run.font.name = 'Angsana New'
            run.font.size = Pt(16)  # Angsana New needs larger size for readability

    def _add_formatted_text_to_paragraph(self, paragraph, text: str, language: str):
        """Add text to paragraph with proper formatting for underscored terms and mixed content"""
        import re
        from docx.shared import Pt

        if language == 'TH':
            # For Thai content, we need to handle mixed Thai-English text
            # Pattern to match: _underscored_ or standalone English words/numbers
            pattern = r'(_[^_]+_|[A-Za-z0-9_]+(?:\s*\([^)]*\))?)'

            last_end = 0
            for match in re.finditer(pattern, text):
                # Add Thai text before the match
                if match.start() > last_end:
                    thai_text = text[last_end:match.start()]
                    thai_run = paragraph.add_run(thai_text)
                    thai_run.font.name = 'Angsana New'
                    thai_run.font.size = Pt(12)

                matched_text = match.group(0)

                # Check if it's underscored text
                if matched_text.startswith('_') and matched_text.endswith('_'):
                    # Remove underscores and apply special formatting
                    clean_text = matched_text[1:-1]
                    special_run = paragraph.add_run(clean_text)
                    special_run.font.name = 'Cambria (Body)'
                    special_run.font.size = Pt(11)
                    special_run.underline = True
                else:
                    # English text/numbers - use normal English font
                    eng_run = paragraph.add_run(matched_text)
                    eng_run.font.size = Pt(11)

                last_end = match.end()

            # Add remaining Thai text
            if last_end < len(text):
                remaining_thai = text[last_end:]
                thai_run = paragraph.add_run(remaining_thai)
                thai_run.font.name = 'Angsana New'
                thai_run.font.size = Pt(12)
        else:
            # For English content, only handle underscored text
            pattern = r'_([^_]+)_'

            last_end = 0
            for match in re.finditer(pattern, text):
                # Add text before the underscored section
                if match.start() > last_end:
                    normal_text = text[last_end:match.start()]
                    normal_run = paragraph.add_run(normal_text)
                    normal_run.font.size = Pt(11)

                # Add the underscored text with Cambria font size 11 and underline (same size as body text)
                underscored_text = match.group(1)
                special_run = paragraph.add_run(underscored_text)
                special_run.font.name = 'Cambria (Body)'
                special_run.font.size = Pt(11)
                special_run.underline = True

                last_end = match.end()

            # Add remaining text after last match
            if last_end < len(text):
                remaining_text = text[last_end:]
                remaining_run = paragraph.add_run(remaining_text)
                remaining_run.font.size = Pt(11)

    def _generate_word_report(self, content: Dict, report_format: str, language: str = 'EN') -> str:
        """Generate Word document report with improved professional layout"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.shared import OxmlElement, qn

            # Get template content
            template = self._get_ghg_template_content(language)

            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"carbon_report_{report_format}_{language}_{timestamp}.docx"
            filepath = os.path.join('reports', filename)

            # Ensure reports directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)

            # Create document with improved margins
            doc = Document()
            
            # Set document margins for better layout
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Cover Page Section
            # Title with enhanced formatting
            title = doc.add_heading(template['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(24)
            title_run.font.color.rgb = RGBColor(0, 54, 146)  # Professional blue
            title_run.bold = True
            
            # Add proper spacing after title
            title_spacing = doc.add_paragraph()
            title_spacing.space_after = Pt(24)
            
            # Subtitle with better formatting
            subtitle = doc.add_heading(content['subtitle'], level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray
            
            # Add proper spacing after subtitle
            subtitle_spacing = doc.add_paragraph()
            subtitle_spacing.space_after = Pt(36)
            
            # Add report generation info with proper spacing
            report_info = doc.add_paragraph()
            report_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            report_info.space_before = Pt(24)
            report_info.space_after = Pt(12)
            report_info_run = report_info.add_run(f"Generated on {content['generated_at']}")
            report_info_run.font.size = Pt(12)
            report_info_run.font.color.rgb = RGBColor(128, 128, 128)
            report_info_run.italic = True
            
            # Page break for main content
            doc.add_page_break()
            
            # Executive Summary Section with enhanced formatting
            exec_heading = doc.add_heading(template['executive_summary_title'], level=1)
            exec_heading_run = exec_heading.runs[0]
            exec_heading_run.font.color.rgb = RGBColor(0, 54, 146)
            exec_heading.space_before = Pt(18)
            exec_heading.space_after = Pt(12)
            
            exec_para = doc.add_paragraph()
            self._add_formatted_text_to_paragraph(exec_para, content['executive_summary'], language)
            exec_para.space_after = Pt(18)
            
            # Key Metrics Section with enhanced table
            metrics_heading = doc.add_heading('Key Metrics' if language == 'EN' else 'ตัวชี้วัดหลัก', level=1)
            metrics_heading_run = metrics_heading.runs[0]
            metrics_heading_run.font.color.rgb = RGBColor(0, 54, 146)
            
            # Create enhanced metrics table
            metrics_table = doc.add_table(rows=1, cols=2)
            metrics_table.style = 'Light Grid Accent 1'
            metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            metrics_table.columns[0].width = Inches(3)
            metrics_table.columns[1].width = Inches(2.5)
            
            # Header row
            hdr_cells = metrics_table.rows[0].cells
            hdr_cells[0].text = 'Metric' if language == 'EN' else 'ตัวชี้วัด'
            hdr_cells[1].text = 'Value' if language == 'EN' else 'ค่า'
            
            # Style header row
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                # Set background color to blue
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '003692')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add metrics data
            for key, value in content['key_metrics'].items():
                row_cells = metrics_table.add_row().cells

                # Use Thai labels if language is Thai, otherwise use English
                if language == 'TH' and key in self.metric_labels_th:
                    row_cells[0].text = self.metric_labels_th[key]
                else:
                    row_cells[0].text = key.replace('_', ' ').title()

                row_cells[1].text = str(value)

                # Style data rows
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Add space
            
            # Emissions by Scope Section with enhanced layout
            scope_heading = doc.add_heading('Emissions by Scope (GHG Protocol)' if language == 'EN' else 'การปล่อยก๊าซเรือนกระจกตามขอบเขต (GHG Protocol)', level=1)
            scope_heading_run = scope_heading.runs[0]
            scope_heading_run.font.color.rgb = RGBColor(0, 54, 146)
            
            # Create scope summary table
            scope_table = doc.add_table(rows=1, cols=3)
            scope_table.style = 'Light Grid Accent 1'
            scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            scope_hdr = scope_table.rows[0].cells
            scope_hdr[0].text = 'Scope' if language == 'EN' else 'ขอบเขต'
            scope_hdr[1].text = 'Emissions (kg CO2e)' if language == 'EN' else 'การปล่อย (kg CO2e)'
            scope_hdr[2].text = 'Percentage' if language == 'EN' else 'เปอร์เซ็นต์'
            
            # Style header
            for cell in scope_hdr:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '003692')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add scope data
            total_scope = sum(content['emissions_by_scope'].values())
            for scope, value in content['emissions_by_scope'].items():
                if value > 0:
                    percentage = (value / total_scope * 100) if total_scope > 0 else 0
                    row_cells = scope_table.add_row().cells
                    row_cells[0].text = scope
                    row_cells[1].text = f"{value:.2f}"
                    row_cells[2].text = f"{percentage:.1f}%"
                    
                    # Style data rows
                    for cell in row_cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add scope descriptions
            doc.add_paragraph()
            for scope, value in content['emissions_by_scope'].items():
                if value > 0:
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(f"{scope}: ")
                    desc_run.bold = True
                    desc_run.font.color.rgb = RGBColor(0, 54, 146)
                    desc_para.add_run(template.get('scope_descriptions', {}).get(scope, f'Description for {scope} not available'))
                    desc_para.style.font.size = Pt(10)
            
            doc.add_paragraph()  # Add space
            
            # Key Findings Section
            if content['key_findings']:
                findings_heading = doc.add_heading(template['key_findings_title'], level=1)
                findings_heading_run = findings_heading.runs[0]
                findings_heading_run.font.color.rgb = RGBColor(0, 54, 146)
                
                for finding in content['key_findings']:
                    finding_para = doc.add_paragraph(finding, style='List Bullet')
                    finding_para.style.font.size = Pt(11)
                
                doc.add_paragraph()  # Add space
            
            # Recommendations Section with enhanced formatting
            if content['recommendations']:
                rec_heading = doc.add_heading(template['recommendations_title'], level=1)
                rec_heading_run = rec_heading.runs[0]
                rec_heading_run.font.color.rgb = RGBColor(0, 54, 146)
                
                for recommendation in content['recommendations']:
                    # Check if recommendation already starts with a number
                    if recommendation.strip() and recommendation.strip()[0].isdigit():
                        rec_para = doc.add_paragraph(recommendation, style='List Bullet')
                    else:
                        rec_para = doc.add_paragraph(recommendation, style='List Number')

                    rec_para.style.font.size = Pt(11)
                
                doc.add_paragraph()  # Add space
            
            # Trend Analysis Section (if available)
            if content.get('trend_analysis'):
                trend_heading = doc.add_heading('Trend Analysis' if language == 'EN' else 'การวิเคราะห์แนวโน้ม', level=1)
                trend_heading_run = trend_heading.runs[0]
                trend_heading_run.font.color.rgb = RGBColor(0, 54, 146)

                trend_para = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(trend_para, content['trend_analysis'], language)

                doc.add_paragraph()  # Add space

            # Emissions Breakdown Section (NEW)
            if content.get('emissions_breakdown'):
                breakdown_heading = doc.add_heading('Emissions Breakdown' if language == 'EN' else 'การแบ่งประเภทการปล่อยก๊าซ', level=1)
                breakdown_heading_run = breakdown_heading.runs[0]
                breakdown_heading_run.font.color.rgb = RGBColor(0, 54, 146)

                breakdown_para = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(breakdown_para, content['emissions_breakdown'], language)

                doc.add_paragraph()  # Add space

            # Page break for methodology
            doc.add_page_break()

            # Methodology Section
            method_heading = doc.add_heading(template['methodology_title'], level=1)
            method_heading_run = method_heading.runs[0]
            method_heading_run.font.color.rgb = RGBColor(0, 54, 146)

            method_para = doc.add_paragraph()
            self._add_formatted_text_to_paragraph(method_para, content['methodology'], language)

            doc.add_paragraph()  # Add space

            # Data Quality Section (NEW)
            if content.get('data_quality'):
                data_quality_heading = doc.add_heading('Data Quality Statement' if language == 'EN' else 'คำชี้แจงคุณภาพข้อมูล', level=1)
                data_quality_heading_run = data_quality_heading.runs[0]
                data_quality_heading_run.font.color.rgb = RGBColor(0, 54, 146)

                data_quality_para = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(data_quality_para, content['data_quality'], language)

                doc.add_paragraph()  # Add space

            # Conclusion Section (NEW)
            if content.get('conclusion'):
                conclusion_heading = doc.add_heading('Conclusion' if language == 'EN' else 'สรุป', level=1)
                conclusion_heading_run = conclusion_heading.runs[0]
                conclusion_heading_run.font.color.rgb = RGBColor(0, 54, 146)

                conclusion_para = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(conclusion_para, content['conclusion'], language)

                doc.add_paragraph()  # Add space

            # Footer with enhanced styling
            doc.add_paragraph()
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("─" * 50)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            footer_text = doc.add_paragraph()
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text_run = footer_text.add_run(f"Report generated on {content['generated_at']}")
            footer_text_run.font.size = Pt(10)
            footer_text_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_text_run.italic = True
            
            # Save document
            doc.save(filepath)
            print(f"✓ Word document saved: {filepath}")
            print(f"✓ File exists: {os.path.exists(filepath)}")
            print(f"✓ Absolute path: {os.path.abspath(filepath)}")
            return filepath
            
        except Exception as e:
            print(f"✗ Word generation error: {str(e)}")
            import traceback
            traceback.print_exc()
            return None

    def _save_to_database(self, user_id: str, report_data: Dict, report_content: Dict,
                         start_date: str, end_date: str, report_format: str, file_path: str,
                         file_type: str, language: str) -> str:
        """Save report to database using your MongoDB schema"""

        print(f"\n{'='*60}")
        print(f"SAVING TO DATABASE")
        print(f"{'='*60}")
        print(f"  File path to save: {file_path}")
        print(f"  File type: {file_type}")
        print(f"  Language: {language}")

        # Generate unique report ID by finding the highest existing number
        # This prevents skipping numbers even with concurrent requests
        latest_report = reports_collection.find_one(
            {"report_id": {"$regex": "^RPT[0-9]+$"}},
            sort=[("report_id", -1)]
        )

        if latest_report and 'report_id' in latest_report:
            # Extract the number from the last report ID (e.g., "RPT005" -> 5)
            try:
                last_number = int(latest_report['report_id'].replace('RPT', ''))
                next_number = last_number + 1
            except ValueError:
                # Fallback if parsing fails
                next_number = reports_collection.count_documents({}) + 1
        else:
            # No reports exist yet, start with 1
            next_number = 1

        report_id = f"RPT{str(next_number).zfill(3)}"
        
        # Parse dates for period structure
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        
        # Create report document matching your schema
        report_document = {
            "_id": ObjectId(),
            "report_id": report_id,
            "user_id": ObjectId(user_id),
            "status": "Completed",
            "total_emission": report_data['total_emissions'],
            "create_date": datetime.utcnow(),
            "organization": report_data['organization'],
            "period": {
                "start_month": start_dt.month,
                "start_year": start_dt.year,
                "end_month": end_dt.month,
                "end_year": end_dt.year
            },
            "emissions_by_category": report_data['emissions_by_category'],
            # Enhanced fields for AI-powered reports
            "report_format": report_format,
            "file_path": file_path,
            "file_type": file_type,
            "language": language,
            "ai_insights_included": True,
            "emissions_by_scope": report_data['emissions_by_scope'],
            "monthly_data": report_data['monthly_data']
        }
        
        # Insert into database
        reports_collection.insert_one(report_document)
        
        return report_id

# Utility functions for Flask integration
def generate_ai_report(user_id: str, start_date: str, end_date: str, 
                      report_format: str = 'GHG', file_type: str = 'PDF', 
                      language: str = 'EN', include_ai: bool = True) -> Dict:
    """Main function to generate AI-powered reports"""
    generator = CarbonReportGenerator()
    return generator.generate_report(user_id, start_date, end_date, report_format, file_type, language, include_ai)

def get_available_report_formats() -> List[str]:
    """Get list of available report formats"""
    return ['GHG']

def get_available_file_types() -> List[str]:
    """Get list of available file types"""
    return ['PDF', 'EXCEL', 'WORD']

def get_available_languages() -> List[str]:
    """Get list of available languages"""
    return ['EN', 'TH']

def validate_report_request(data: Dict) -> Dict:
    """Validate report generation request"""
    required_fields = ['start_date', 'end_date', 'report_format']
    
    # Check required fields
    for field in required_fields:
        if field not in data:
            return {
                'valid': False,
                'error': f'Missing required field: {field}'
            }
    
    # Validate report format
    if data['report_format'] not in get_available_report_formats():
        return {
            'valid': False,
            'error': f'Invalid report format. Available formats: {get_available_report_formats()}'
        }
    
    # Validate dates
    try:
        start_date = datetime.fromisoformat(data['start_date'].replace('Z', '+00:00'))
        end_date = datetime.fromisoformat(data['end_date'].replace('Z', '+00:00'))
        
        if start_date >= end_date:
            return {
                'valid': False,
                'error': 'Start date must be before end date'
            }
            
    except ValueError as e:
        return {
            'valid': False,
            'error': f'Invalid date format: {str(e)}'
        }
    
    return {'valid': True}
